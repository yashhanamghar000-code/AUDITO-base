import os
import concurrent.futures
from functools import partial
from typing import List, Optional

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pypdf import PdfReader
import pdfplumber
import pytesseract
import fitz  # PyMuPDF
from PIL import Image
import docx

# Your i5-1235U has 2 Performance + 8 Efficient cores = 10 logical cores.
# The Celery worker now runs up to 3 files concurrently (concurrency=3, see
# docker-compose.yml) instead of 1, so each file's internal thread count is
# capped lower here to compensate — 3 files x 3 threads = 9, still leaves
# headroom for the FastAPI process, the embedding model, and the OS.
# Override with an env var if you move this to a bigger machine.
DEFAULT_PARSER_WORKERS = max(2, min(3, (os.cpu_count() or 4) - 2))
PARSER_MAX_WORKERS = int(os.getenv("PARSER_MAX_WORKERS", DEFAULT_PARSER_WORKERS))


class MultiUserParser:
    @classmethod
    def parse_uploaded_stream(
        cls,
        file_path: str,
        file_name: str,
        user_id: str,
        session_id: str,
        file_id: str,
    ) -> List[Document]:
        """
        Main multi-tenant entry point. Reads from a file on shared disk
        (written by the FastAPI upload endpoint, read by the Celery
        worker) rather than passing raw bytes through the message broker.

        `file_id` is the Postgres UploadedFile.id (as a string), created
        BEFORE this task runs (see main.py's /api/upload). It gets stamped
        onto every chunk's metadata below so that later, a specific upload
        can be deleted or selected-for-search even if another file shares
        the exact same filename.
        """
        ext = os.path.splitext(file_name)[1].lower()
        print(f"\n[Parser Gateway] User: {user_id} | Session: {session_id} | File: {file_id} | Processing: {file_name}")

        if ext == ".pdf":
            documents = cls._parse_pdf(file_path, file_name, user_id, session_id, file_id)
        elif ext in [".docx", ".doc"]:
            documents = cls._parse_docx(file_path, file_name, user_id, session_id, file_id)
        elif ext in [".txt", ".csv", ".md"]:
            documents = cls._parse_text_file(file_path, file_name, user_id, session_id, file_id)
        elif ext in [".png", ".jpg", ".jpeg", ".tiff"]:
            documents = cls._parse_image_ocr(file_path, file_name, user_id, session_id, file_id)
        else:
            print(f"Unsupported file format: {ext}")
            return []

        if not documents:
            print("No text could be extracted from the file.")
            return []

        return cls._chunk_documents(documents, file_name, user_id, session_id)

    # ------------------------------------------------------------------
    # PDF — parallelized across CPU cores, one thread per page batch.
    # Each worker re-opens the file from disk by path (file handles/state
    # can't be safely shared across threads once we start mutating page
    # rotation), which is why every helper below opens its own handles
    # instead of sharing one from the caller.
    # ------------------------------------------------------------------
    @staticmethod
    def _parse_page_worker(file_path: str, file_name: str, user_id: str, session_id: str, file_id: str, page_num: int) -> Optional[Document]:
        try:
            tables = []
            raw_text = ""
            image_summary_context = ""
            used_fallback = False
            needs_forced_rotation = False
            has_images = False

            # --- STEP 1: pdfplumber pass — layout inspection + normal extraction ---
            with pdfplumber.open(file_path) as pdf:
                if page_num > len(pdf.pages):
                    return None
                page = pdf.pages[page_num - 1]

                # --- GLOBAL PROTECTION FOR PDFPLUMBER LAYOUT CRASHES ---
                try:
                    has_images = len(page.images) > 0 or len(page.rects) > 15
                    if has_images:
                        image_summary_context = "\n[Visual Content Note: Page contains embedded visual layers.]"

                    # Quick orientation check — landscape-shaped page is often
                    # a rotated portrait page (common in financial-report scans/exports).
                    if page.width > page.height:
                        needs_forced_rotation = True
                    else:
                        chars = page.chars
                        if chars and len(chars) > 100:
                            sampled_chars = chars[::20]
                            vertical_chars = sum(
                                1 for c in sampled_chars
                                if c.get("orientation") in ["up", "down"] or c.get("upright") == 0
                            )
                            if vertical_chars / len(sampled_chars) > 0.30:
                                needs_forced_rotation = True
                except Exception:
                    # If pdfplumber trips on a layout/index error, skip straight
                    # to the pypdf rotation-normalized path below.
                    needs_forced_rotation = False
                    used_fallback = True

                # --- STEP 2: NATIVE ROTATION CHECK + CORRECTION VIA PYPDF ---
                # NOTE: we open a fresh PdfReader here (rather than sharing one
                # across threads) because pypdf_page.rotate() mutates page
                # state. Sharing a single reader across ThreadPoolExecutor
                # workers risks one thread's rotation call bleeding into
                # another thread's page. The extra open() is cheap relative
                # to getting garbled/reversed text on rotated pages.
                pypdf_page = None
                native_rotation = 0
                try:
                    local_reader = PdfReader(file_path)
                    pypdf_page = local_reader.pages[page_num - 1]
                    native_rotation = pypdf_page.get("/Rotate", 0)
                except Exception:
                    print(f"   ⚠️ Page {page_num}: Could not fetch page rotation metadata.")

                if (native_rotation in [90, 270] or needs_forced_rotation) and pypdf_page is not None and not used_fallback:
                    rotation_angle = (360 - native_rotation) if native_rotation in [90, 270] else 90
                    print(f"   -> Page {page_num}: Adjusting layout by {rotation_angle}° to normalize horizontal reading axis...")
                    try:
                        pypdf_page.rotate(rotation_angle)
                        raw_text = pypdf_page.extract_text() or ""
                        used_fallback = True
                    except Exception as e:
                        print(f"   ⚠️ Rotation parsing stream bottleneck on Page {page_num}: {e}")

                # --- STEP 3: Normal extraction path if rotation wasn't needed ---
                if not used_fallback:
                    try:
                        tables = page.extract_tables()
                        raw_text = page.extract_text() or ""
                    except Exception:
                        print(f"   ⚠️ Layout stream bottleneck on Page {page_num}. Executing recovery...")
                        try:
                            if pypdf_page is not None:
                                raw_text = pypdf_page.extract_text() or ""
                                used_fallback = True
                        except Exception:
                            print(f"   ❌ Critical Error: Page {page_num} unreadable. Skipping.")
                            return None

                # --- STEP 4: last-resort recovery if fallback path produced nothing ---
                if used_fallback and not raw_text and pypdf_page is not None:
                    try:
                        raw_text = pypdf_page.extract_text() or ""
                    except Exception:
                        print(f"   ❌ Critical Recovery Error: Page {page_num} completely unparseable.")
                        return None

            # --- STEP 5: OCR fallback only when extracted text is genuinely poor/missing ---
            # This is the expensive path, so most pages skip it entirely.
            if len(raw_text.strip()) < 50:
                try:
                    with fitz.open(file_path) as fitz_doc:
                        fitz_page = fitz_doc.load_page(page_num - 1)
                        zoom = 150 / 72
                        pix = fitz_page.get_pixmap(matrix=fitz.Matrix(zoom, zoom))
                        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                        ocr_text = pytesseract.image_to_string(img, config="--psm 6")
                        if len(ocr_text.strip()) > len(raw_text.strip()):
                            raw_text = ocr_text
                            image_summary_context += "\n[System Note: Text extracted via OCR.]"
                        img.close()
                except Exception as e:
                    print(f"OCR failed on page {page_num}: {e}")

            cleaned_lines = [line.strip() for line in raw_text.splitlines() if line.strip()]
            sanitized_text = "\n".join(cleaned_lines)

            table_markdown = ""
            if tables and not used_fallback:
                for table in tables:
                    cleaned_table = [[str(cell) if cell is not None else "" for cell in row] for row in table]
                    for row in cleaned_table:
                        if any(row):
                            table_markdown += "| " + " | ".join(row) + " |\n"
                    table_markdown += "\n"

            combined_content = f"ATTENTION LLM: FILE: '{file_name}' | PAGE: {page_num} {image_summary_context}\n" + sanitized_text
            if table_markdown:
                combined_content += "\n\n### Extracted Document Tables:\n" + table_markdown
            elif used_fallback:
                combined_content += "\n\n[System Note: Text structurally layout-normalized and parsed via horizontal fallback stream.]"

            return Document(
                page_content=combined_content,
                metadata={
                    "source": file_name,
                    "page": page_num,
                    "has_table": bool(table_markdown),
                    "has_images": has_images,
                    "was_rotated": used_fallback,
                    "user_id": user_id,
                    "session_id": session_id,
                    "file_id": file_id,
                },
            )
        except Exception as e:
            print(f"Error processing page {page_num}: {e}")
            return None

    @classmethod
    def _parse_pdf(cls, file_path: str, file_name: str, user_id: str, session_id: str, file_id: str) -> List[Document]:
        try:
            reader = PdfReader(file_path)
            total_pages = len(reader.pages)
        except Exception as e:
            print(f"Failed to read PDF pages: {e}")
            return []

        if total_pages == 0:
            return []

        print(f"[Thread Parser] {total_pages} pages across up to {PARSER_MAX_WORKERS} thread workers...")

        # Small documents: skip pool overhead entirely for a 1-2 page file
        if total_pages <= 2 or PARSER_MAX_WORKERS <= 1:
            results = [
                cls._parse_page_worker(file_path, file_name, user_id, session_id, file_id, p)
                for p in range(1, total_pages + 1)
            ]
            return [r for r in results if r is not None]

        # Properly aligned variables for the multi-threaded code execution path
        worker_fn = partial(cls._parse_page_worker, file_path, file_name, user_id, session_id, file_id)
        page_numbers = list(range(1, total_pages + 1))
        parent_documents: List[Document] = []

        # ThreadPoolExecutor sidesteps Windows multiprocessing serialization completely
        with concurrent.futures.ThreadPoolExecutor(max_workers=PARSER_MAX_WORKERS) as executor:
            for res in executor.map(worker_fn, page_numbers):
                if res is not None:
                    parent_documents.append(res)

        # executor.map preserves entry order chronology seamlessly
        return parent_documents

    # ------------------------------------------------------------------
    # Other formats — small enough that per-page parallelism doesn't matter.
    # ------------------------------------------------------------------
    @staticmethod
    def _parse_docx(file_path: str, file_name: str, user_id: str, session_id: str, file_id: str) -> List[Document]:
        doc = docx.Document(file_path)
        full_text = [para.text for para in doc.paragraphs if para.text.strip()]
        content = "\n".join(full_text)

        return [Document(
            page_content=f"ATTENTION LLM: FILE: '{file_name}'\n" + content,
            metadata={"source": file_name, "page": 1, "user_id": user_id, "session_id": session_id, "file_id": file_id},
        )]

    @staticmethod
    def _parse_text_file(file_path: str, file_name: str, user_id: str, session_id: str, file_id: str) -> List[Document]:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()
        return [Document(
            page_content=f"ATTENTION LLM: FILE: '{file_name}'\n" + content,
            metadata={"source": file_name, "page": 1, "user_id": user_id, "session_id": session_id, "file_id": file_id},
        )]

    @staticmethod
    def _parse_image_ocr(file_path: str, file_name: str, user_id: str, session_id: str, file_id: str) -> List[Document]:
        try:
            img = Image.open(file_path)
            text = pytesseract.image_to_string(img)
            img.close()
        except Exception as e:
            print(f"OCR execution failed on image file: {e}")
            text = ""

        return [Document(
            page_content=f"ATTENTION LLM: IMAGE FILE: '{file_name}'\n[System Note: Text extracted via OCR]\n" + text,
            metadata={"source": file_name, "page": 1, "is_image": True, "user_id": user_id, "session_id": session_id, "file_id": file_id},
        )]

    @classmethod
    def _chunk_documents(cls, parent_documents: List[Document], original_file_name: str, user_id: str, session_id: str) -> List[Document]:
        final_chunks = []
        SAFE_PAGE_CHAR_LIMIT = 3000
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=2500, chunk_overlap=300, length_function=len)

        for doc in parent_documents:
            if len(doc.page_content) <= SAFE_PAGE_CHAR_LIMIT:
                final_chunks.append(doc)
                continue

            if "### Extracted Document Tables:" in doc.page_content:
                narrative, table_block = doc.page_content.split("### Extracted Document Tables:", 1)
                table_block = "### Extracted Document Tables:" + table_block
            else:
                narrative, table_block = doc.page_content, ""

            sub_docs = text_splitter.split_documents([Document(page_content=narrative, metadata=doc.metadata)])
            for sub in sub_docs:
                if table_block:
                    sub.page_content += "\n\n" + table_block
                final_chunks.append(sub)

        os.makedirs("./debug_logs", exist_ok=True)
        debug_output_path = f"./debug_logs/audit_{user_id}_{session_id}.txt"

        with open(debug_output_path, "a", encoding="utf-8") as txt_file:
            txt_file.write(f"=== CHUNK AUDIT LOG: {original_file_name} | {len(final_chunks)} SEGMENTS ===\n\n")
            for idx, chunk in enumerate(final_chunks, start=1):
                txt_file.write(f"--- CHUNK {idx} | Source: {chunk.metadata['source']} | Page: {chunk.metadata['page']} ---\n")
                txt_file.write(chunk.page_content)
                txt_file.write("\n\n" + "=" * 50 + "\n\n")

        return final_chunks